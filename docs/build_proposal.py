# -*- coding: utf-8 -*-
"""2026 ICT 표준 챌린지 공모전 데모 개발 기획서 생성 스크립트.
실행: python docs/build_proposal.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "맑은 고딕"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.6  # 160%
BLACK = RGBColor(0, 0, 0)

doc = Document()

style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = FONT_SIZE
style.font.color.rgb = BLACK
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), FONT_NAME)
style.paragraph_format.line_spacing = LINE_SPACING
style.paragraph_format.space_after = Pt(0)

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(0.8)


def set_run_font(run, bold=False, size=None, color=BLACK):
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_para(target, text="", bold=False, size=None, align=None, space_after=3):
    p = target.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, bold=bold, size=size)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def tighten_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            row.cells[idx].width = Cm(w)


def label_cell(cell, text):
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, bold=True, size=Pt(11))
    shade_cell(cell, "FCE7C8")
    tighten_cell_margins(cell)


def content_cell(cell, lines):
    cell.paragraphs[0].text = ""
    tighten_cell_margins(cell)
    first = True
    for item in lines:
        text, bold = item if isinstance(item, tuple) else (item, False)
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        set_run_font(r, bold=bold, size=Pt(10.5))


# ============================================================
# 표지 (간략)
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("2026 ICT 표준 챌린지 공모전 데모 개발 기획서")
set_run_font(r, bold=True, size=Pt(16))
title_p.paragraph_format.space_after = Pt(2)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("출품작: FactON(팩트온) · 제출: FactON팀")
set_run_font(r, size=Pt(10))
sub_p.paragraph_format.space_after = Pt(8)

# ============================================================
# 표 1
# ============================================================
table1 = doc.add_table(rows=5, cols=2)
table1.style = "Table Grid"
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(table1, [2.6, 15.4])

label_cell(table1.rows[0].cells[0], "출품작\n제목")
content_cell(table1.rows[0].cells[1], [
    ("정부 공식반박 우선 연결 기반 AI 근거검증 서비스 \u2018FactON(팩트온)\u2019", True),
])

label_cell(table1.rows[1].cells[0], "제품·서비스\n유형")
content_cell(table1.rows[1].cells[1], [
    "웹 서비스(Next.js 반응형 웹앱, PC·모바일 브라우저 즉시 이용)",
    "SW데모 구동 영상·소스코드 첨부: ☑ 예   ☐ 아니오",
])

label_cell(table1.rows[2].cells[0], "출품작\n요약")
content_cell(table1.rows[2].cells[1], [
    "SNS·카톡 속 주장을 붙여넣으면 AI가 핵심 주장을 추출해 정부의 \u2018사실은 이렇습니다\u2019 "
    "공식반박과 우선 매칭하고, 매칭되지 않으면 공개 웹 근거를 다중 AI로 교차검증합니다. "
    "AI가 진위를 직접 판정하는 대신 정부 공식자료를 우선 연결해 신뢰도를 높인 "
    "시민용 사실확인 서비스입니다.",
])

label_cell(table1.rows[3].cells[0], "SW개발\n아이디어의\n참신성")
content_cell(table1.rows[3].cells[1], [
    ("배경: ", True),
    "2026.7.7 개정 정보통신망법 시행(허위조작정보 반복 유통 시 최대 10억원 과징금) 등 "
    "AI 생성정보 신뢰 문제가 실제 법적 이슈가 됐고, 정치인·유명인 사칭 영상으로 인한 "
    "투자사기 피해도 발생 중입니다. 그런데 기존 AI 팩트체크 다수는 LLM이 스스로 진위를 "
    "판정해 환각·정치적 편향이라는 또 다른 신뢰 문제를 만듭니다.",
    ("차별점: ", True),
    "FactON은 반대로 1차는 정부가 이미 반박한 사례와 매칭해 AI의 새 판정이 아닌 정부 "
    "원문을 그대로 연결(고신뢰)하고, 매칭 실패 시에만 2차로 AI GMS 1차\u2192AI GMS 2차\u2192"
    "로컬 AI GMS 순차 폴백의 웹검색 교차검증을 거쳐 근거확인·상충·불충분 3단계로 "
    "표시합니다. 근거 불충분 시 거짓으로 단정하지 않고 안전하게 보류합니다.",
    ("표준 조합의 새로움: ", True),
    "TTA 표준 4종을 입력품질\u2192주장·의미정확성\u2192신뢰성·설명가능성\u2192개인정보보호라는 "
    "파이프라인 4단계에 각각 매핑해 결합 적용했습니다. 사용자는 \u2018AI 판단\u2019이 아닌 "
    "\u2018정부 반박 원문\u2019을 직접 보게 되어 정치적 중립성 논란을 최소화합니다.",
])

label_cell(table1.rows[4].cells[0], "TTA 표준\n활용성 및\n문제해결\n연결성")
c = table1.rows[4].cells[1]
c.paragraphs[0].text = ""
tighten_cell_margins(c)
p = c.paragraphs[0]
p.paragraph_format.space_after = Pt(3)
r = p.add_run("< 본 제품·서비스 개발에 활용한 TTA 표준 >")
set_run_font(r, bold=True, size=Pt(10.5))

inner = c.add_table(rows=5, cols=3)
inner.style = "Table Grid"
set_col_widths(inner, [0.9, 4.3, 10.8])
hdr = inner.rows[0].cells
for i, t in enumerate(["NO", "TTA 표준 번호", "TTA 표준명"]):
    hdr[i].paragraphs[0].text = ""
    tighten_cell_margins(hdr[i])
    rp = hdr[i].paragraphs[0].add_run(t)
    set_run_font(rp, bold=True, size=Pt(10))
    shade_cell(hdr[i], "E5E7EB")

std_rows = [
    ("1", "TTAK.KO-10.1344-Part2", "유통·활용 데이터 점검 방법 - 제2부: 비정형 데이터 품질지표"),
    ("2", "TTAK.KO-10.1419", "한국어 음성 및 텍스트 데이터의 의미적 정확성 품질검증 방법"),
    ("3", "TTAK.KO-10.1497", "인공지능 시스템 신뢰성 제고를 위한 요구사항"),
    ("4", "TTAK.KO-12.0414", "인공지능(AI) 서비스 개인정보보호 프레임워크"),
]
for i, (no, code, name) in enumerate(std_rows, start=1):
    cells = inner.rows[i].cells
    for j, val in enumerate([no, code, name]):
        cells[j].paragraphs[0].text = ""
        tighten_cell_margins(cells[j])
        rp = cells[j].paragraphs[0].add_run(val)
        set_run_font(rp, size=Pt(10))

lines = [
    ("① TTAK.KO-10.1344-Part2 \u2014 ", True),
    "checkInputQuality.ts에서 입력 텍스트 길이·품질을 확인해 분석 가능 여부를 판단"
    "(원문 6.3, 7.4 반영).",
    ("② TTAK.KO-10.1419 \u2014 ", True),
    "extractClaims.ts·verifyClaim.ts에서 주장을 문장 단위로 분리하고 근거와의 연결을 "
    "확인, 답을 못 찾으면 거짓 단정 대신 \u2018근거 불충분\u2019 처리(원문 5, 6.4 반영).",
    ("③ TTAK.KO-10.1497 \u2014 ", True),
    "route.ts의 diagnostics와 화면의 \u2018검증 품질 체크\u2019 패널에서 각 단계 성공·실패를 "
    "판정 이유·원문 링크와 함께 표시하고 실패 시 안전하게 보류(REQ.04·05·11·13·14 반영).",
    ("④ TTAK.KO-12.0414 \u2014 ", True),
    "maskPII.ts에서 전화번호·주민등록번호를 마스킹하고 입력 원문은 서버에 미저장"
    "(원문 7.3.1~7.3.5 반영).",
]
for item in lines:
    text, bold = item if isinstance(item, tuple) else (item, False)
    pp = c.add_paragraph()
    pp.paragraph_format.line_spacing = LINE_SPACING
    pp.paragraph_format.space_after = Pt(2)
    rp = pp.add_run(text)
    set_run_font(rp, bold=bold, size=Pt(10))

# ============================================================
# 다이어그램 (작게, 표 2 앞에 인라인 배치 - 페이지 브레이크 없음)
# ============================================================
add_para(doc, "< FactON 검증 파이프라인 및 TTA 표준 4종 매핑 >", bold=True, size=Pt(10.5), space_after=2)
pic_p = doc.add_paragraph()
pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_p.paragraph_format.space_after = Pt(4)
run = pic_p.add_run()
run.add_picture(r"C:\Users\SSAFY\Desktop\TTA\facton\docs\pipeline_diagram.png", width=Cm(13.5))

# ============================================================
# 표 2
# ============================================================
table2 = doc.add_table(rows=4, cols=2)
table2.style = "Table Grid"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(table2, [2.6, 15.4])

label_cell(table2.rows[0].cells[0], "생성형 AI\n활용성")
content_cell(table2.rows[0].cells[1], [
    ("① AI GMS 1차 검증: ", True),
    "extractClaims.ts로 핵심 주장 1~5개를 추출하고, verifyWithClaudeSearch.ts는 "
    "web_search 도구로 정부·공공기관·원자료를 우선 검색해 수치·기관·시점 단위로 "
    "비교, 근거확인/상충/불충분을 판정합니다.",
    ("② AI GMS 2차 폴백: ", True),
    "AI GMS 1차 실패(쿼터 소진·장애) 시 Google 검색 그라운딩 또는 검색결과 재판정으로 "
    "자동 전환됩니다(verifyWithGeminiSearch.ts, verifyWithGemmaEvidence.ts).",
    ("③ 로컬 AI GMS(경량 모델) \u2014 3차 폴백: ", True),
    "위 두 단계가 모두 실패해도 로컬 LLM이 근거를 비교·판정해 서비스 연속성을 "
    "보장합니다(verifyClaim.ts).",
    ("④ 개발 과정 활용: ", True),
    "생성형 AI 코딩 도구로 요구사항 분석부터 구현, 실제 운영 버그(폴백 매칭 오탐, 타임아웃, "
    "API 쿼터 초과) 진단·수정까지 전 과정을 수행했습니다. 개발 중 실제 "
    "API 쿼터 소진 상황에서 AI GMS 2차\u2192로컬 AI GMS로 자동 전환되는 것을 재현 검증했습니다.",
])

label_cell(table2.rows[1].cells[0], "SW개발\n구현완성도·\n서비스가능성")
content_cell(table2.rows[1].cells[1], [
    ("핵심기능: ", True),
    "1) 텍스트/URL\u2192주장 추출\u2192정책브리핑 \u2018사실확인\u2019 API 1차 매칭(실연동) "
    "2) 매칭 실패 시 AI GMS 1차\u2192AI GMS 2차\u2192로컬 AI GMS 웹검색 교차검증\u2192Fact Card 3단계 "
    "표시(문제구절·교정문장·원문링크) 3) 결과 공유(ShareCard)와 부가 화면(/how, "
    "/literacy 자가진단, /report 신고안내, /sources 출처 18종 디렉토리)",
    ("실제 작동: ", True),
    "정책뉴스 API 실연동으로 최근 180일 정책자료 약 2,400여건을 조회·캐싱함을 실측.",
    ("오류 대응: ", True),
    "외부 서비스 장애 시 단계적 폴백으로 무중단 운영, 실제 장애(크레딧 소진·속도제한·"
    "타임아웃) 상황에서 정상 폴백을 재현 테스트로 확인.",
    ("보완 필요: ", True),
    "키워드 겹침 폴백 매칭 정밀도 고도화 진행 중, 이미지·영상 입력(OCR/STT)은 후순위.",
])

label_cell(table2.rows[2].cells[0], "공공편익·\n포용성")
content_cell(table2.rows[2].cells[1], [
    "카톡 가족방 등 고령층에 집중되는 생활정보 허위정보 피해를 예방하고, 흩어진 정부 "
    "공식자료 접근성을 높여 정보 비대칭을 해소합니다.",
    "\u2018AI가 임의로 단정하지 않는\u2019 설계로 정치적 중립성을 확보합니다.",
    "/literacy 자가진단, /report 신고 안내(KISO 연계), /sources 출처 디렉토리로 시민의 "
    "정보 판별 습관 형성과 접근성을 함께 지원합니다.",
    "결과를 가족 단톡방에 바로 공유하는 카드(ShareCard)로 재확산 방지까지 유도합니다.",
])

label_cell(table2.rows[3].cells[0], "기타")
content_cell(table2.rows[3].cells[1], [
    ("가산점 \u2014 공공데이터 사용: ", True),
    "☑ 예   ☐ 아니오",
    "\u2013 데이터명: 정책브리핑 정책뉴스 서비스(policyNewsList2) / URL: "
    "apis.data.go.kr/1371000/policyNewsService2/policyNewsList2 / 제공기관: "
    "문화체육관광부·한국정책방송원",
    ("기타: ", True),
    "커밋을 요구사항분석\u2192스키마설계\u21921차매칭\u21922차fallback\u2192판정보류엔진\u2192UI\u2192"
    "개인정보마스킹\u2192테스트\u2192리팩터링 단계로 분리, 각 커밋에 생성형 AI 활용 방식을 "
    "기록했습니다. 참고: TTA 표준 4종 원문, 정책브리핑 정책뉴스 API 명세.",
])

# OOXML은 표로 문서를 끝낼 수 없어 표 뒤에 문단이 하나 필요한데, 표2가 페이지 하단에
# 거의 닿아 있으면 이 문단 하나 때문에 빈 5페이지가 추가로 생긴다. 폰트를 최소화해
# 이 문단이 표 바로 아래 여백에 들어가도록 한다.
tail_p = doc.add_paragraph()
tail_p.paragraph_format.space_before = Pt(0)
tail_p.paragraph_format.space_after = Pt(0)
tail_p.paragraph_format.line_spacing = 1.0
tail_run = tail_p.add_run("")
set_run_font(tail_run, size=Pt(1))
# 빈 문단의 높이는 문단 마크(pilcrow) 자체의 rPr이 결정하고, 방금 넣은 run의 크기와는
# 무관하다 - 문단 마크 rPr도 똑같이 줄여야 실제로 한 줄 높이가 줄어든다.
tail_pPr = tail_p._p.get_or_add_pPr()
mark_rpr = OxmlElement("w:rPr")
mark_sz = OxmlElement("w:sz")
mark_sz.set(qn("w:val"), "2")
mark_rpr.append(mark_sz)
tail_pPr.append(mark_rpr)

out_path = r"C:\Users\SSAFY\Desktop\TTA\facton\2026 ICT 표준 챌린지 공모전 기획서_FactON팀.docx"
doc.save(out_path)
print("saved:", out_path)
